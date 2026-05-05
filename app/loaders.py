"""
loaders.py – načítavanie vstupných dokumentov pre FMEA pipeline.

Podporované formáty:
- .txt, .md  – obyčajný text
- .pdf       – cez pypdf (extrakcia textu, bez OCR)
- .docx      – paragrafy aj tabuľky (kontrolné plány, pracovné inštrukcie)
- .xlsx      – všetky hárky cez openpyxl (kontrolné plány, existujúce FMEA tabuľky)

Chyby pri načítaní jedného súboru sú zalogované cez `logging`, ostatné
súbory sa spracujú aj tak. Word lock files (~$xxx.docx) a skryté súbory
sa preskakujú.
"""

from pathlib import Path
import logging

from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


# ── Loadery jednotlivých formátov ──────────────────────────────────────────────

def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)

    return "\n".join(pages)


def load_docx(path: Path) -> str:
    """
    Extrahuje text z DOCX – paragrafy aj obsah tabuliek.

    Pôvodná verzia iterovala iba `doc.paragraphs`, čím úplne ignorovala
    obsah tabuliek. Pritom kontrolné plány, pracovné inštrukcie a procesné
    toky sú v praxi často štruktúrované práve do tabuliek a relevantný
    kontext by sa stratil ešte pred odoslaním do AI.

    Riadky tabuľky sú serializované ako tab-separované hodnoty, čo zachová
    aspoň základný stĺpcový vzťah pre tokenizáciu v AI modeli.
    """
    doc = Document(str(path))
    parts: list[str] = []

    # 1) Paragrafy v hlavnom toku dokumentu
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # 2) Tabuľky – každý riadok ako tab-separovaný text
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            non_empty = [c for c in cells if c]
            if non_empty:
                parts.append("\t".join(non_empty))

    return "\n".join(parts)


def load_xlsx(path: Path) -> str:
    """
    Extrahuje text zo všetkých hárkov XLSX súboru.

    Každý hárok je v texte oddelený nadpisom s jeho názvom, aby AI mohla
    rozlíšiť kontext jednotlivých listov (napr. "Control Plan", "FMEA",
    "Procesný tok"). Vzorce sú vyhodnotené pomocou `data_only=True`,
    čo vráti naposledy uloženú vypočítanú hodnotu namiesto reťazca
    so vzorcom.

    `read_only=True` je dôležité pre veľké kontrolné plány s tisíckami
    riadkov – knižnica openpyxl by inak nahrala celý súbor do pamäte.
    """
    wb = load_workbook(filename=str(path), data_only=True, read_only=True)
    parts: list[str] = []

    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"# Hárok: {sheet_name}")

            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(value).strip()
                    for value in row
                    if value is not None and str(value).strip()
                ]
                if cells:
                    parts.append("\t".join(cells))
    finally:
        wb.close()

    return "\n".join(parts)


# ── Dispatcher ────────────────────────────────────────────────────────────────

# Registrácia loaderov podľa prípony. Pridanie nového formátu znamená
# pridať jeden riadok – nemusíme rozširovať reťaz if/elif v load_documents.
LOADERS = {
    ".txt":  load_txt,
    ".md":   load_txt,
    ".pdf":  load_pdf,
    ".docx": load_docx,
    ".xlsx": load_xlsx,
}


def load_documents(folder: str) -> list[dict]:
    """
    Prejde rekurzívne `folder` a načíta všetky podporované dokumenty.

    Vracia zoznam dictov v tvare {"source": str, "text": str}. Súbory
    neznámych formátov sa ignorujú, chyby pri načítaní sú zalogované
    bez prerušenia spracovania ostatných dokumentov.
    """
    docs: list[dict] = []

    for path in Path(folder).rglob("*"):
        if not path.is_file():
            continue

        # Preskočiť Word/Excel lock files (~$xxx.docx) a skryté súbory.
        if path.name.startswith("~$") or path.name.startswith("."):
            continue

        loader = LOADERS.get(path.suffix.lower())
        if loader is None:
            continue

        try:
            text = loader(path)
        except Exception as exc:
            logger.error("Chyba pri načítaní %s: %s", path, exc)
            continue

        if text.strip():
            docs.append({
                "source": str(path),
                "text":   text,
            })
        else:
            logger.warning("Súbor %s je prázdny – preskakujem.", path)

    return docs 